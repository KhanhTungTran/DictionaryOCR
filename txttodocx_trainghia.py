from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Inches, RGBColor
import json

def searchForUnderline(para, text, start, end, indexes, bold, italic):
    cuts = []
    for index in indexes:
        if start <= index and index < end:
            cuts.append(index - start)
        else:   break
    # print(cuts, ' ', text)

    low = 0
    for index in cuts:
        high = index
        # if high > 0:
        for i in range(low, high):
            run = para.add_run(text[i])
            run.bold = bold
            run.italic = italic
            if text[i] == '*':
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        run = para.add_run(text[index])
        run.bold = bold
        run.italic = italic
        run.underline = True
        if text[index] == '*':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        low = high + 1

    for i in range(low, end - start):
        run = para.add_run(text[i])
        run.bold = bold
        run.italic = italic
        if text[i] == '*':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

inputsDir = 'results'
path = os.listdir(inputsDir)
# Đi từng dòng, dòng nào là '' thì bỏ qua
# Tìm những cụm như ' ph.', ' vch.', ... => quay ngược lại tìm dấu chấm trước đó, in đậm từ dấu chấm đến trước cụm,
# và thêm '\n' trước dấu chấm vừa tìm được
doc = Document()
para = ''
wordTypes = { 'bt.': 'biến từ', 'chđt.': 'chỉ định từ', 'dt.': 'danh từ', 'đdt.': 'đại danh từ', 'đt.': 'động từ', 'gt.': 'giới từ', 'lt.': 'liên từ', 'pht.': 'phó từ', 'st.': 'số từ', 'tt.': 'tĩnh từ', 'trt.': 'trạng từ', 'tht.': 'thán từ', 'vt.': 'vấn từ'}

docPara = doc.add_paragraph('')
docPara.add_run('KÍ HIỆU:').bold = True

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('*').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
docPara.add_run(': các kí tự lạ nằm ngoài bảng chữ cái Tiếng Việt')

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('_').bold = True
docPara.add_run(': các chữ cái có xác suất dự đoán sai cao')

docPara = doc.add_paragraph('\n')
docPara = doc.add_paragraph('')


for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    # par = doc.add_paragraph(txtFile)
    contents = txtFile.split('\n')
    cumTu = False
    vidu = False
    first = False
    capTu = False
    for line in contents:
        if line == '':
            continue
        
        indexes = []
        # print(line[line.rfind('['):line.rfind(']') + 1])
        if line[line.rfind('['):line.rfind(']') + 1] != '':
            indexes = json.loads(line[line.rfind('['):line.rfind(']') + 1]) # string to list
            indexes = sorted(indexes)

        line = line[:line.rfind('[')]

        if line.startswith('Ví dụ'):
            docPara = doc.add_paragraph('')
            paraFormat = docPara.paragraph_format
            paraFormat.left_indent = Inches(0.5)
            searchForUnderline(docPara, line, 0, len(line), indexes, False, False)
            # docPara.add_run(line)
            capTu = False

        elif line.find(' - ') != -1:
            if line.isupper():
                docPara = doc.add_paragraph('')
                searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                # docPara.add_run(line).bold = True
                docPara.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                capTu = False
            elif line[:line.find(' - ')].isupper():
                docPara = doc.add_paragraph('')
                searchForUnderline(docPara, line[:line.find(' - ')], 0, line.find(' - '), indexes, True, False)
                searchForUnderline(docPara, line[line.find(' - '):], line.find(' - '), len(line), indexes, False, False)
                # docPara.add_run(line[:line.find(' - ')]).bold = True
                # docPara.add_run(line[line.find(' - '):])
                capTu = False
            elif capTu == True:
                docPara = doc.add_paragraph('')
                searchForUnderline(docPara, line, 0, len(line), indexes, True, True)
                # run = docPara.add_run(line)
                # run.font.bold = True
                # run.font.italic = True
                docPara.paragraph_format.left_indent = Inches(0.5)      

        elif line == 'Gặp từ trái nghĩa:' or line == 'Cặp từ trái nghĩa:':
            docPara = doc.add_paragraph('')
            paraFormat = docPara.paragraph_format
            paraFormat.left_indent = Inches(0.5)
            run = docPara.add_run('Cặp từ trái nghĩa:')
            run.font.bold = True
            run.font.underline = True
            capTu = True

        elif capTu == True:
            docPara = doc.add_paragraph('')
            searchForUnderline(docPara, line, 0, len(line), indexes, False, False)
            # run = docPara.add_run(line)
            docPara.paragraph_format.left_indent = Inches(0.5)  

        else:
            for i in range(len(line)):
                run = docPara.add_run(line[i])
                if line[i] == '*':
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                
doc.save('result.docx')