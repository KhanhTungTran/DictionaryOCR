from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Inches

inputsDir = 'results'
path = os.listdir(inputsDir)
path = sorted(path, key = lambda x: len(x))

# doc = Document()
# para = ''
# docPara = doc.add_paragraph('')

# tags = {' d.': ' danh từ', 'đg.': 'động từ', ' t.': ' tính từ', 'đ.': 'đại từ', ' p.': ' phụ từ', 'k.': 'kết từ', 'tr.': 'trợ từ', ' c.': ' cảm từ', 'hoài nghi đt.': 'hoài nghi động từ'}

# notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'} # còn nữa

# tags = { 'cd.': 'ca dao', 'dt.': 'danh từ', 'đt.': 'động từ', 'gt.': 'giới từ', 'id.': 'ít dùng', 'lt.': 'liên từ', ' ng.': ' nghĩa', 'pt.': 'phụ từ', 'tht.': 'thán từ', 'tng.': 'tục ngữ', 'trt.': 'trợ từ', 'vt.': 'vị từ'}

# notes = {'(id.).': '(ít dùng)', '(kng.)': '(khẩu ngữ)' , '(thgt.)': '(thông tục)', '(ph.)': '(phương ngữ)', '(vchg.)': '(văn chương)',  '(trtr.)': '(trang trọng)', '(kc.)': '(kiểu cách)', '(chm.)': '(chuyên môn)'} # còn nữa

tags = {' dt.': ' danh từ', 'đgt.': 'động từ', ' tt.': ' tính từ', ' pht.': ' phụ từ',}

notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'}

alphabets = {'a': ['a', 'ă', 'â', 'à', 'á', 'ã', 'ả', 'ạ', 'ắ', 'ằ', 'ẵ', 'ẳ', 'ặ', 'ấ', 'ầ', 'ẫ', 'ẩ', 'ậ'], 'b': ['b'], 'c': ['c'], 'd': ['d', 'đ'], 'e': ['e', 'ê', 'é', 'è', 'ẽ', 'ẻ', 'ẹ', 'ế', 'ề', 'ễ', 'ể', 'ệ'], 'f': ['f'], 'g': ['g'], 'h': ['h'], 'i': ['i', 'í', 'ì', 'ĩ', 'ỉ', 'ị'], 'j': ['j'], 'k': ['k'], 'l': ['l'], 'm': ['m'], 'n': ['n'], 'o': ['o', 'ô', 'ơ', 'ó', 'ò', 'õ', 'ỏ', 'ọ', 'ố', 'ồ', 'ỗ', 'ổ', 'ộ', 'ớ', 'ờ', 'ỡ', 'ở', 'ợ'], 'p': ['p'], 'q': ['q'], 'r': ['r'], 's': ['s'], 't': ['t'], 'u': ['u', 'ư', 'ú', 'ù', 'ũ', 'ủ', 'ụ', 'ứ', 'ừ', 'ữ', 'ử', 'ự'], 'v': ['v'], 'w': ['w'], 'x': ['x'], 'y': ['y'], 'z': ['z']}

extrasBeforeTag = {'cn.': '(cũng nói)', 'cv.': 'cũng viết'}  # tiếp theo sẽ là 1 từ in nghiêng có chấm cuối câu 
# extrasAfterTag = {'x.': 'xem' }
extrasAfterTag = {' x.': ' xem'}
newEntry = True
currentAlphabet = '`'
lastLine = '.'

words = []
types = []
meanings = []
exs = []

def hasSubString(text, tags):
    for tag in tags.keys():
        if tag in text:
            return True
    return False

def hasExtraTag(text):
    for tag in extrasAfterTag.keys():
        if tag in text:
            index = text.find(tag)
            if index > 1 and text[index - 1] != '.' or text[index - 1] != '!' or text[index - 1] != '?' or text[index - 1] != ',':
                return True

    return False

def checkFirstWord(text, currAlphabet, alphabetsDict):
    for char in alphabetsDict[currAlphabet]:
        if text.lower().startswith(char) or text.lower().startswith('"' + char):
            return True
    return False

def checkOrderNumber(text):
    for ordNum in ['II', 'III', 'IV', 'V']:
        if text.startswith(ordNum):
            return ordNum
    return ''

for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    txtFile = txtFile.replace('Ä', 'A')
    txtFile = txtFile.replace('ä', 'a')
    txtFile = txtFile.replace('“', '"')
    txtFile = txtFile.replace('”', '"')
    txtFile = txtFile.replace('¿', 't')
    txtFile = txtFile.replace(':.', 't.')
    txtFile = txtFile.replace('¡(', 't')

    for key, value in notes.items():
        txtFile = txtFile.replace(key, value)

    for key, value in extrasBeforeTag.items():
        txtFile = txtFile.replace(key, value)
        
    contents = txtFile.split('\n')
    
    for line in contents:
        if line == '':
            newEntry = True
            continue

        # if line.lower() == chr(ord(currentAlphabet) + 1) + chr(ord(currentAlphabet) + 1):
        #     continue
        
        if newEntry or line.lower().startswith(chr(ord(currentAlphabet) + 1) + ',' + chr(ord(currentAlphabet) + 1)) or ((lastLine.endswith('.') or lastLine.endswith('!') or lastLine.endswith('?') or lastLine.endswith(',')) and (checkFirstWord(line, currentAlphabet, alphabets) or checkOrderNumber(line))):
            # print(line)
            newEntry = False
            nextAlphabet = chr(ord(currentAlphabet) + 1)
            if line.lower().startswith(nextAlphabet + ',' + nextAlphabet):
                currentAlphabet = nextAlphabet
                words.append(nextAlphabet + ',' + nextAlphabet.upper())
                meanings.append(line[4:])
                types.append('Ý nghĩa chữ cái')

            # else:
            elif (lastLine.endswith('.') or lastLine.endswith('!') or lastLine.endswith('?') or lastLine.endswith(',')):
                if (hasSubString(line.replace(',', '.').lower(), tags) or hasExtraTag(line.replace(',', '.'))): # mục từ mới
                    if line.lower().startswith(chr(ord(currentAlphabet) + 1)) or line.lower().startswith('"' + chr(ord(currentAlphabet) + 1)):
                        currentAlphabet = chr(ord(currentAlphabet) + 1)
                    
                    elif (not checkFirstWord(line, currentAlphabet, alphabets)) and checkOrderNumber(line) == '':
                        if lastLine.endswith(' '):
                            meanings[-1] += line
                        else:
                            meanings[-1] += ' ' + line
                        continue
                    # elif not (line.lower().startswith(currentAlphabet) or line.lower().startswith('"' + currentAlphabet)):
                    #     yes = False
                    #     for ordNum in ['II', 'III', 'IV', 'V']:
                    #         if line.startswith(ordNum):
                    #             yes = True
                    #             break
                    #     if not yes:
                    #         continue

                    currentTags = {}
                    
                    tempLine = line.replace(',', '.').lower()
                    for extra in extrasAfterTag.keys():
                        if extra in line:
                            currentTags[extra] = line.find(extra)
                    for tag in tags.keys():
                        if tag in tempLine:
                            currentTags[tag] = tempLine.find(tag)
                    sortedCurrentTags = sorted(currentTags.items(), key=lambda kv: kv[1])
                    word = line[:sortedCurrentTags[0][1]]
                    content = line[sortedCurrentTags[len(sortedCurrentTags) - 1][1] + len(sortedCurrentTags[len(sortedCurrentTags) - 1][0]):]

                    words.append(word)

                    if sortedCurrentTags[0][0] in tags.keys():
                        types.append(tags[sortedCurrentTags[0][0]])
                    else:
                        types.append(extrasAfterTag[sortedCurrentTags[0][0]])

                    for tagKey,_ in sortedCurrentTags[1:]:
                        if tagKey in tags.keys():
                            if tagKey.startswith(' '):
                                types[-1] += tags[tagKey]
                            else:
                                types[-1] += ' ' + tags[tagKey]
                        else:
                            types[-1] += ' ' + extrasAfterTag[tagKey]
                    meanings.append(content)

                elif checkFirstWord(line, currentAlphabet, alphabets):
                    for i in range(2, len(line)):
                        if line[i].isupper() and line[i - 1] == ' ' and (line[i - 2] != '.' and line[i - 2] != ',' and line[i - 2] != '!' and line[i - 2] != '?'):
                            words.append(line[:i])
                            meanings.append(line[i:])
                            types.append('')
                            break
                    else:
                        # docPara.add_run(line)
                        meanings[-1] += line
                elif checkOrderNumber(line) != '':
                    num = checkOrderNumber(line)
                    words.append(num)
                    meanings.append(line[len(num):])
                    types.append('')
                else:
                    words.append('')
                    meanings.append(line)
                    types.append('')
        elif lastLine.endswith(' '):
            # docPara.add_run(line)
            meanings[-1] += line
        else:
            # docPara.add_run(' ' + line)
            meanings[-1] += ' ' + line
        lastLine = line

for i in range(len(words)):
    if len(words[i]) >= 1 and words[i][0] == ' ':
        words[i] = words[i][1:]
    if len(words[i]) >= 1 and words[i][-1] == ' ':
        words[i] = words[i][:-1]

    if len(meanings[i]) >= 1 and meanings[i][0] == ' ':
        meanings[i] = meanings[i][1:]
    if len(meanings[i]) >= 1 and meanings[i][-1] == ' ':
        meanings[i] = meanings[i][:-1]

    if len(types[i]) >= 1 and types[i][0] == ' ':
        types[i] = types[i][1:]
    if len(types[i]) >= 1 and types[i][-1] == ' ':
        types[i] = types[i][:-1]

import xml.etree.ElementTree as ET
import json
import json

def indent(elem, level=0):
    i = "\n" + level*"  "
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = i + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
        for elem in elem:
            indent(elem, level+1)
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = i

tree = ET.parse('result.xml')
root = tree.getroot()
for i in range(len(words)):
    element = root.makeelement('MUC_TU', {'Noi_dung': words[i], 'Loai_tu': types[i]})
    root.append(element)
    ET.SubElement(root[i], 'Y_NGHIA', {'Noi_dung': meanings[i]})

indent(root)
tree.write('result.xml', encoding='utf-8', xml_declaration=True)
