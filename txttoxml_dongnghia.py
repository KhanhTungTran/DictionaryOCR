from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import RGBColor
import json

def searchForUnderline(text, start, end, indexes, bold, italic):
    res = ''
    cuts = []
    for index in indexes:
        if start <= index and index < end:
            cuts.append(index - start)
        else:   break
    # print(cuts, ' ', text)

    low = 0
    for index in cuts:
        high = index
        # if high > 0:
        for i in range(low, high):
            res += text[i]
            # run = para.add_run(text[i])
            # run.bold = bold
            # run.italic = italic
            # if text[i] == '*':
            #     run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        res += '[' + text[index] + ']'
        # run = para.add_run(text[index])
        # run.bold = bold
        # run.italic = italic
        # run.underline = True
        # if text[index] == '*':
        #     run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        low = high + 1

    for i in range(low, end - start):
        res += text[i]
        # run = para.add_run(text[i])
        # run.bold = bold
        # run.italic = italic
        # if text[i] == '*':
        #     run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    return res

inputsDir = 'results'
path = os.listdir(inputsDir)

words = []
groups = []
contents = []
exs = []

for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    # par = doc.add_paragraph(txtFile)
    content = txtFile.split('\n')
    cumTu = False
    first = False
    vidu = False
    for line in content:
        if line == '':
            continue

        indexes = []
        # print(line[line.rfind('['):line.rfind(']') + 1])
        if line[line.rfind('['):line.rfind(']') + 1] != '':
            # print(line)
            indexes = json.loads(line[line.rfind('['):line.rfind(']') + 1]) # string to list 
            indexes = sorted(indexes)
            
        line = line[:line.rfind('[')]
        # print(line)
        if line.find(' - ') != -1:
            cumTu = False
            vidu = False
            # docPara = doc.add_paragraph('')
            if line.lower().find('ví dụ') != -1:
                # print(line)
                vidu = True
                words.append(searchForUnderline(line[:line.find(' - ')], 0, line.find(' - '), indexes, True, False))
                while len(exs) < len(words):
                    exs.append('')
                while len(words) > len(groups):
                    groups.append(groups[len(groups) - 1])
                contents.append(searchForUnderline(line[line.find(' - ') + 3:line.lower().find('ví dụ')], line.find(' - ') + 3, line.lower().find('ví dụ'), indexes, False, False))
                # exs.append('')
                while line.lower().find('ví dụ') != -1 and line.lower().find(':') != -1:
                    exs[len(exs) - 1] += (searchForUnderline(line[line.lower().find('ví dụ'):line.find(':') + 1], line.lower().find('ví dụ'), line.find(':') + 1, indexes, False, True))
                    line = line[line.find(':') + 1:]
                    # print(line)
                    if line.lower().find('ví dụ') != -1 and line.lower().find(':') != -1:
                        exs[len(exs) - 1] += searchForUnderline(line[0:line.lower().find('ví dụ')], 0, line.lower().find('ví dụ'), indexes, False, False)
                    else:
                        exs[len(exs) - 1] += searchForUnderline(line, 0, len(line), indexes, False, False)
            else:
                words.append(searchForUnderline(line[:line.find(' - ')], 0, line.find(' - '), indexes, True, False))
                while len(exs) < len(words):
                    exs.append('')
                while len(words) > len(groups):
                    groups.append(groups[len(groups) - 1])
                contents.append(searchForUnderline(line[line.find(' - ') + 3:], line.find(' - ') + 3, len(line), indexes, False, False))

        elif line.lower().find('ví dụ') != -1:
            vidu = True
            # print(line)

            contents[len(contents) - 1] += searchForUnderline(line[:line.find('ví dụ')], 0, line.find('ví dụ'), indexes, False, True)
            exs.append(searchForUnderline(line[line.find(':'):], line.find(':'), len(line), indexes, False, False))

        elif cumTu == True and not line.isupper():
            if not first:
                pass
                # docPara.add_run(' ').bold = True
                # searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
            else:
                first = False
                pass
                # docPara = doc.add_paragraph('')
                # searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                # docParaFormat = docPara.paragraph_format
                # docParaFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                

        elif line.isupper():
            if cumTu == True:
                pass
                # docPara.add_run(' ').bold = True
                # searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                # docPara.add_run(' ' + line).bold = True
            else:
                groups.append(searchForUnderline(line, 0, len(line), indexes, True, False))
                vidu = False
                # docPara = doc.add_paragraph('')
                # searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                # docParaFormat = docPara.paragraph_format
                # docParaFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cumTu = True
            first = True
        
        else:
            # docPara.add_run(' ')
            # searchForUnderline(docPara, line, 0, len(line), indexes, False, False)
            if vidu:
                exs[len(exs) - 1] += searchForUnderline(line, 0, len(line), indexes, False, False)
            else:
                contents[len(contents) - 1] += searchForUnderline(line, 0, len(line), indexes, False, False)
# exs.append('')
# groups.append(groups[len(groups) - 1])

import xml.etree.ElementTree as ET

def indent(elem, level=0):
    i = "\n" + level*"  "
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = i + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
        for elem in elem:
            indent(elem, level+1)
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = i

tree = ET.parse('result.xml')
root = tree.getroot()
for i in range(len(contents)):
    element = root.makeelement('MUC_TU', {'Noi_dung': words[i], 'Loai_tu': groups[i]})
    root.append(element)
    ET.SubElement(root[i], 'Y_NGHIA', {'Noi_dung': contents[i], 'Minh_hoa': exs[i]})

indent(root)
tree.write('result.xml', encoding='utf-8', xml_declaration=True)
