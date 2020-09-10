import cv2
from os import listdir
from os.path import isfile, join
import pytesseract
import numpy as np
import pandas as pd

import sys
import os
import time
import argparse

import torch
import torch.nn as nn
import torch.backends.cudnn as cudnn
from torch.autograd import Variable

from PIL import Image

import cv2
from skimage import io
import numpy as np
from craft_utils import adjustResultCoordinates, getDetBoxes, getDetBoxes_core, getPoly_core
from imgproc import *
from file_utils import *
import json
import zipfile

from craft import CRAFT
from collections import OrderedDict

from HoughLine import splitImageToColumns

text_threshold = 0.7
low_text = 0.4
link_threshold = 0.4
# cuda = True
cuda = False
canvas_size = 1280
mag_ratio = 1.5
trained_model_path = './craft_mlt_25k.pth'
poly = False
refine = False
show_time = False
refine_net = None

inputDir = 'inputs'
imageName = list(filter(lambda file: file[-3:] == 'jpg', os.listdir(inputDir)))
columnDir = 'splitColumn'
resultsDir = 'results'

# splitImageToColumns(imageName, inputDir, columnDir)
for image in imageName:
    print('---------------------',image)
    img = cv2.imread(inputDir + '/' + image)
    (H, W) = img.shape[:2]
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    ret, thes = cv2.threshold(gray,0,255,cv2.THRESH_BINARY+cv2.THRESH_OTSU)

    colList = []

    try:
        # Crop column
        h, w = img.shape[:2]
        dst = cv2.erode(thes, kernel=np.ones((30, 10)))
        hist = cv2.reduce(dst, 0, cv2.REDUCE_AVG).reshape(-1)
        th = min(list(filter(lambda x: x > 245 and x < 255, hist)))
        uppers = [y for y in range(w-1) if hist[y]>th and hist[y+1]<=th]

        lowers = [y for y in range(w-1) if hist[y]<=th and hist[y+1]>th]
        # print('lowers :', lowers)
        potentialPair = []
        for iU in range(len(uppers)):
            for iL in range(len(lowers)):
                if (uppers[iU] < lowers[iL] and abs(uppers[iU] - lowers[iL]) > w / 4 and abs(uppers[iU] - lowers[iL]) < w / 2):
                    potentialPair.append([iU, iL])
        leftCols = list(filter(lambda x: uppers[x[0]] < w / 3, potentialPair))
        rightCols = list(filter(lambda x: lowers[x[1]] > 2 * w / 3, potentialPair))
    except Exception as e:
        print(image)
        continue
    
    if (len(leftCols) == 0 or len(rightCols) == 0):
        print("File " + image + " error! Please recheck it.")
        continue
    else:
        leftCol = min(leftCols, key=lambda x: lowers[x[1]] - uppers[x[0]])
        rightCol = min(rightCols, key=lambda x: lowers[x[1]] - uppers[x[0]])
        expand = round((lowers[leftCol[1]] - uppers[leftCol[0]]) * 0.03)
        colList.append(thes[0:h,uppers[leftCol[0]] - expand:lowers[leftCol[1]] + expand])
        expand = round((lowers[rightCol[1]] - uppers[rightCol[0]]) * 0.03)
        colList.append(thes[0:h,uppers[rightCol[0]] - expand:lowers[rightCol[1]] + expand])

    # Skew correction for each column
    for i, col in enumerate(colList):
        try:
            gray = cv2.bitwise_not(col)
            pts = cv2.findNonZero(gray)
            ret = cv2.minAreaRect(pts)
            (cx,cy), (w,h), angle = ret
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            M = cv2.getRotationMatrix2D((cx,cy), angle, 1.0)
            rotated = cv2.warpAffine(col, M, (col.shape[1], col.shape[0]), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            cv2.imwrite(columnDir + '/' + image[0:-4] + '-' + str(i) + '.jpg', rotated)
        except Exception as e:
            print(image)
            continue


def copyStateDict(state_dict):
    if list(state_dict.keys())[0].startswith("module"):
        start_idx = 1
    else:
        start_idx = 0
    new_state_dict = OrderedDict()
    for k, v in state_dict.items():
        name = ".".join(k.split(".")[start_idx:])
        new_state_dict[name] = v
    return new_state_dict


def str2bool(v):
    return v.lower() in ("yes", "y", "true", "t", "1")


def test_net(net, image, text_threshold, link_threshold, low_text, cuda, poly, refine_net=None):
    t0 = time.time()

    # resize
    img_resized, target_ratio, size_heatmap = resize_aspect_ratio(image, canvas_size, interpolation=cv2.INTER_LINEAR,
                                                                  mag_ratio=mag_ratio)
    ratio_h = ratio_w = 1 / target_ratio

    # preprocessing
    x = normalizeMeanVariance(img_resized)
    x = torch.from_numpy(x).permute(2, 0, 1)  # [h, w, c] to [c, h, w]
    x = Variable(x.unsqueeze(0))  # [c, h, w] to [b, c, h, w]
    if cuda:
        x = x.cuda()

    # forward pass
    with torch.no_grad():
        y, feature = net(x)

    # make score and link map
    score_text = y[0, :, :, 0].cpu().data.numpy()
    score_link = y[0, :, :, 1].cpu().data.numpy()

    # refine link
    if refine_net is not None:
        with torch.no_grad():
            y_refiner = refine_net(y, feature)
        score_link = y_refiner[0, :, :, 0].cpu().data.numpy()

    t0 = time.time() - t0
    t1 = time.time()

    # Post-processing
    boxes, polys = getDetBoxes(score_text, score_link, text_threshold, link_threshold, low_text, poly)

    # coordinate adjustment
    boxes = adjustResultCoordinates(boxes, ratio_w, ratio_h)
    polys = adjustResultCoordinates(polys, ratio_w, ratio_h)
    for k in range(len(polys)):
        if polys[k] is None: polys[k] = boxes[k]

    t1 = time.time() - t1

    # render results (optional)
    render_img = score_text.copy()
    render_img = np.hstack((render_img, score_link))
    ret_score_text = cvt2HeatmapImg(render_img)

    if show_time: print("\ninfer/postproc time : {:.3f}/{:.3f}".format(t0, t1))

    return boxes, polys, ret_score_text


net = CRAFT()
net.load_state_dict(copyStateDict(torch.load(trained_model_path, map_location='cpu')))
net.eval()

import copy

UNCLASSIFIED = -2
NOISE = -1


class Point:
    def __init__(self, x, y, id):
        self.x = x
        self.y = y
        self.id = id
        self.cluster_id = UNCLASSIFIED

    def __repr__(self):
        return '(x:{}, y:{}, id:{}, cluster:{})' \
            .format(self.x, self.y, self.id, self.cluster_id)


def n_pred(p1, p2):
    return (p1.x - p2.x) ** 2 / 160000 + (p1.y - p2.y) ** 2 / 225 <= 1


def w_card(points):
    return len(points)


def GDBSCAN(points, n_pred, min_card, w_card):
    points = copy.deepcopy(points)
    cluster_id = 0
    for point in points:
        if point.cluster_id == UNCLASSIFIED:
            if _expand_cluster(points, point, cluster_id, n_pred, min_card,w_card):
                cluster_id = cluster_id + 1
    clusters = {}
    for point in points:
        key = point.cluster_id
        if key in clusters:
            clusters[key].append(point)
        else:
            clusters[key] = [point]
    return list(clusters.values())


def _expand_cluster(points, point, cluster_id, n_pred, min_card, w_card):
    if not _in_selection(w_card, point):
        points.change_cluster_id(point, UNCLASSIFIED)
        return False

    seeds = points.neighborhood(point, n_pred)
    if not _core_point(w_card, min_card, seeds):
        points.change_cluster_id(point, NOISE)
        return False

    points.change_cluster_ids(seeds, cluster_id)
    seeds.remove(point)

    while len(seeds) > 0:
        current_point = seeds[0]
        result = points.neighborhood(current_point, n_pred)
        if w_card(result) >= min_card:
            for p in result:
                if w_card([p]) > 0 and p.cluster_id in [UNCLASSIFIED, NOISE]:
                    if p.cluster_id == UNCLASSIFIED:
                        seeds.append(p)
                    points.change_cluster_id(p, cluster_id)
        seeds.remove(current_point)
    return True


def _in_selection(w_card, point):
    return w_card([point]) > 0


def _core_point(w_card, min_card, points):
    return w_card(points) >= min_card


class Points:
    def __init__(self, points):
        self.points = points

    def __iter__(self):
        for point in self.points:
            yield point

    def __repr__(self):
        return str(self.points)

    def get(self, index):
        return self.points[index]

    def neighborhood(self, point, n_pred):
        return list(filter(lambda x: n_pred(point, x), self.points))

    def change_cluster_ids(self, points, value):
        for point in points:
            self.change_cluster_id(point, value)

    def change_cluster_id(self, point, value):
        index = (self.points).index(point)
        self.points[index].cluster_id = value

    def labels(self):
        return set(map(lambda x: x.cluster_id, self.points))

def imgtotext(pic):
    # croped = cv2.imread(file_path, cv2.IMREAD_GRAYSCALE)
    croped = cv2.cvtColor(pic, cv2.COLOR_BGR2GRAY)

    x = y = 0
    h = croped.shape[0]
    w = croped.shape[1]

    mask = np.ones(croped.shape, np.uint8) * 255
    mask[h - int(h * 0.9):int(h * 0.9), x:x + w] = croped[h - int(h * 0.9):int(h * 0.9), x:x + w]
    # cv2.imwrite('check1.jpg', croped)
    if (croped.shape[0] > 150):
        scale_digit = 150 / croped.shape[0]
        croped = cv2.resize(mask, (int(mask.shape[1] * scale_digit), int(mask.shape[0] * scale_digit)),
                            cv2.INTER_LANCZOS4)
        # croped = cv2.resize(croped, (int(croped.shape[1] * scale_digit), int(croped.shape[0] * scale_digit)))
    # croped = cv2.resize(mask, (int(mask.shape[1] * 2), int(mask.shape[0] * 2)), cv2.INTER_LANCZOS4)
    # cv2.imwrite('check2.jpg', croped)
    
    kernel = np.ones((5,5),np.float32)/25
    # croped = cv2.morphologyEx(croped, cv2.MORPH_OPEN, kernel)
    croped = cv2.filter2D(croped,-1,kernel)

    custom_oem_psm_config = '--oem 3 --psm 7 --dpi 300'
    return pytesseract.image_to_string(croped, lang='vie', config=custom_oem_psm_config)


from os import listdir
from os.path import isfile, join

image_index = 0

# files = ['test/' + f for f in listdir('test/') if isfile(join('test/', f))]
files = ['splitColumn/' + f for f in listdir('splitColumn/') if isfile(join('splitColumn/', f))]

# dataset = []

import docx

resultWordDoc = docx.Document()

# Xử lí từng bức hinh đã được tách thành 1 cột trong 1 trang của từ điển:
for file_path in files:
    dataset = []
    print('------------------'+file_path)
    print(image_index)

    image = loadImage(file_path) # image là một numpy array
    # cv2.imwrite('1.jpg',image)
    # bboxes: các bounded box, polys: các polygon
    bboxes, polys, score_text = test_net(net, image, text_threshold, link_threshold, low_text, cuda, poly, refine_net)
    

    poly_indexes = {}
    central_poly_indexes = []
    for i in range(len(polys)):
        poly_indexes[i] = polys[i]
        x_central = (polys[i][0][0] + polys[i][1][0] + polys[i][2][0] + polys[i][3][0]) / 4
        y_central = (polys[i][0][1] + polys[i][1][1] + polys[i][2][1] + polys[i][3][1]) / 4
        central_poly_indexes.append({i: [int(x_central), int(y_central)]})

    X = []
    for idx, x in enumerate(central_poly_indexes):
        point = Point(x[idx][0], x[idx][1], idx)
        X.append(point)
        
    # n_pred: a neighborhood predicate
    # 1: minimum weight
    # w_card: weight function
    clustered = GDBSCAN(Points(X), n_pred, 1, w_card)
    
    # clustered = sorted(clustered, key = lambda elem : (elem.x % 100, elem.y % 100))
    cluster_values = []
    for cluster in clustered:
        sort_cluster = sorted(cluster, key=lambda elem: (elem.x, elem.y))
        max_point_id = sort_cluster[len(sort_cluster) - 1].id
        min_point_id = sort_cluster[0].id
        max_rectangle = sorted(poly_indexes[max_point_id], key=lambda elem: (elem[0], elem[1]))
        min_rectangle = sorted(poly_indexes[min_point_id], key=lambda elem: (elem[0], elem[1]))

        right_above_max_vertex = max_rectangle[len(max_rectangle) - 1]
        right_below_max_vertex = max_rectangle[len(max_rectangle) - 2]
        left_above_min_vertex = min_rectangle[0]
        left_below_min_vertex = min_rectangle[1]
        cluster_values.append([left_above_min_vertex, left_below_min_vertex, right_above_max_vertex, right_below_max_vertex])

    img = np.array(image[:,:,::-1])
    img_2 = img.copy()
    # cv2.imwrite('image.jpg',img_2)
    # cv2.waitKey()
    def compare(elem, other):
        elem_top_left_corner = [0, 0]
        other_top_left_corner = [0, 0]

        elem_top_left_corner[0] = sorted(elem, key = lambda val : val[0])[0][0]
        elem_top_left_corner[1] = sorted(elem, key = lambda val : val[1])[0][1]
        other_top_left_corner[0] = sorted(other, key = lambda val : val[0])[0][0]
        other_top_left_corner[1] = sorted(other, key = lambda val : val[1])[0][1]

        if elem_top_left_corner[1] - other_top_left_corner[1] >= -40.0 and elem_top_left_corner[1] - other_top_left_corner[1] <= 40.0:
            if elem_top_left_corner[0] < other_top_left_corner[0]:
                return -1
            else:
                return 1
        elif elem_top_left_corner[1] - other_top_left_corner[1] > 40.0:
            return 1
        else:
            return -1

    from functools import cmp_to_key

    sorted_cluster_values = sorted(cluster_values, key = cmp_to_key(compare))
    cluster_values = sorted_cluster_values
    for i, box in enumerate(cluster_values):
        poly = np.array(box).astype(np.int32).reshape((-1))
        poly = poly.reshape(-1, 2)
        rect = cv2.boundingRect(poly)
        x,y,w,h = rect

        ##### boxes around text
        # test_img = cv2.rectangle(img_2, (x, y), (x + w, y + h), (0, 0, 255), 2)
        # cv2.imwrite('test_img.jpg',test_img)

        # print(str(x) + '_' + str(y))
        # if h/w < 0.5:
        # croped1 = img[y-5:y+h+5, x-5:x+w+5].copy()
        croped1 = img[abs(abs(y)-10):abs(y)+h+10, abs(abs(x)-10):abs(x)+w+10].copy()
        # cv2.imwrite('croppedimage.jpg',croped1)
        # cv2.waitKey()
        text = imgtotext(croped1)
        try:
            # dataset.append((imgtotext(croped1)))
            dataset.append(text)
            image_index += 1
        except Exception as e:
            break
    # print(dataset)
    poly = False
    refine = False
    show_time = False
    refine_net = None
    # resultWordDoc.add_paragraph(dataset[0])
    # for word in dataset[1:]:
    #     resultWordDoc.add_run(' ' + word)
    with open(resultsDir + '/' + file_path[12:-3] + 'txt', 'w', encoding='utf8') as f:
        for row in dataset:
            f.write(row + '\n') 

    # break
# resultWordDoc.save(resultsDir + '/' + file_path[12:-3] + 'docx')