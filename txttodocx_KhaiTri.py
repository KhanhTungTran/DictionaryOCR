
from docx import Document
import re
import os

inputsDir = 'results'
path = os.listdir(inputsDir)
# path = ['Tu dien viet nam Ban tu thu Khai tri_131-0.txt']

# Đi từng dòng, dòng nào là '' thì bỏ qua
# Tìm những cụm như ' ph.', ' vch.', ... => quay ngược lại tìm dấu chấm trước đó, in đậm từ dấu chấm đến trước cụm,
# và thêm '\n' trước dấu chấm vừa tìm được
doc = Document()
para = ''
wordTypes = { 'bt.': 'biến từ', 'chđt.': 'chỉ định từ', 'dt.': 'danh từ', 'đdt.': 'đại danh từ', 'đt.': 'động từ', 'gt.': 'giới từ', 'lt.': 'liên từ', 'pht.': 'phó từ', 'st.': 'số từ', 'tt.': 'tính từ', 'trt.': 'trạng từ', 'tht.': 'thán từ', 'vt.': 'vấn từ', 'xt.': 'xem tiếng'}
docPara = doc.add_paragraph('')

currentAlphabet = 'a'
countNextAlphabet = 0

for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    # par = doc.add_paragraph(txtFile)
    contents = txtFile.split('\n')
    for line in contents:
        if line == '':
            continue
        i = 1
        while line[:i].isupper() and i <= len(line):
            i += 1
        
        i -= 2
        newWord = False
        wordType = ''
        if i == 1 and len(line) >= 4:
            if line[i + 1] == '(':
                newWord = True
            elif line[i + 1: i + 4] in wordTypes or line[i + 1: i + 5] in wordTypes or line[i + 1: i + 6] in wordTypes:
                newWord = True
        elif i > 1:
            newWord = True
        
        if newWord:
            if line[i + 1: i + 4] in wordTypes:
                wordType = line[i + 1: i + 4]
            elif line[i + 1: i + 5] in wordTypes:
                wordType = line[i + 1: i + 5]
            elif line[i + 1: i + 6] in wordTypes:
                wordType = line[i + 1: i + 6]
            newPara = line[:i]
            line = line[i:]
            # print(line)

            # docPara.add_run(para)
            literals = para.split()
            for literal in literals:
                if literal.replace(',', '.') in wordTypes.keys():
                    docPara.add_run(' ' + wordTypes[literal.replace(',', '.')]).italic = True
                else:
                    docPara.add_run(' ' + literal)

            docPara = doc.add_paragraph('')
            docPara.add_run(newPara).bold = True
            # if wordType != '':
            #     line = line[1 + len(wordType) - 1]
            #     docPara.add_run(' ')
            #     docPara.add_run(wordTypes[wordType]).italic = True
            para = ''
            
        # for wordType in wordTypes.keys():
        #     if wordType in line:
        #         if line.find(wordType) + len(wordType) < len(line) and line[line.find(wordType) + len(wordType)] == ')':
        #             continue
        #         newPara = para[para.rfind('.') + 1:] + line[:line.find(wordType)]
        #         line = line[line.find(wordType) + len(wordType):]
        #         para = para[:para.rfind('.') + 1]
        #         # break
        #         docPara.add_run(para)
        #         docPara = doc.add_paragraph('')
        #         docPara.add_run(newPara).bold = True
        #         # docPara.add_run(' ')
        #         docPara.add_run(wordTypes[wordType]).italic = True
        #         para = ''
        #         break
        
        para += (line + ' ')

doc.save('result.docx')