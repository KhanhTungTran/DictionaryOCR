from docx import Document
from docx.shared import RGBColor
import re
import os
import json

inputsDir = 'results'
path = os.listdir(inputsDir)
# Đi từng dòng, dòng nào là '' thì bỏ qua
# Tìm những cụm như ' ph.', ' vch.', ... => quay ngược lại tìm dấu chấm trước đó, in đậm từ dấu chấm đến trước cụm,
# và thêm '\n' trước dấu chấm vừa tìm được
doc = Document()
para = ''
# wordTypes = {'c.': 'cảm từ', 'd.': 'danh từ', 'đ.': 'đại từ', 'đg.': 'động từ', 'k.': 'kết từ', 'p.': 'phụ từ', 't.': 'tính từ', 'tr.': 'trợ từ', 'x.': 'xem'}
wordTypes = ['cảm từ', 'danh từ', 'động từ', 'đại từ', 'kết từ', 'phụ từ', 'tính từ', 'trợ từ', ' xem ']

docPara = doc.add_paragraph('')
docPara.add_run('KÍ HIỆU:').bold = True

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('*').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
docPara.add_run(': các kí tự lạ nằm ngoài bảng chữ cái Tiếng Việt')

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('_').bold = True
docPara.add_run(': các chữ cái có xác suất dự đoán sai cao')

docPara = doc.add_paragraph('\n')
docPara = doc.add_paragraph('')
for txt in path:
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    # par = doc.add_paragraph(txtFile)
    contents = txtFile.split('\n')
    for line in contents:
        # if line == '':
        #     continue
        # done = False
        indexes = []
        # print(line[line.rfind('['):line.rfind(']') + 1])
        if line[line.rfind('['):line.rfind(']') + 1] != '':
            # print('something')
            indexes = json.loads(line[line.rfind('['):line.rfind(']') + 1]) # string to list

        line = line[:line.rfind('[')]
        bold = -1
        italic = len(line)
        regular = 0

        for wordType in wordTypes:
            if wordType in line:
                wordTypeIndex = line.find(wordType)
                bold = wordTypeIndex - 1
                italic = wordTypeIndex + len(wordType)
                regular = wordTypeIndex + len(wordType)
                # docPara.add_run(line[:wordTypeIndex]).bold = True
                # docPara.add_run(wordType).italic = True
                # docPara.add_run(line[wordTypeIndex + len(wordType):])
                # docPara = doc.add_paragraph('')
                # done = True
                # newPara = para[para.rfind('.') + 1:] + line[:line.find(wordType)]
                # line = line[line.find(wordType) + len(wordType):]
                # para = para[:para.rfind('.') + 1]
                # # break
                # docPara.add_run(para)
                # docPara = doc.add_paragraph('')
                # docPara.add_run(newPara).bold = True
                # docPara.add_run(wordTypes[wordType]).italic = True
                # para = ''
                break
        # if not done:
        #     docPara.add_run(line)
        #     docPara = doc.add_paragraph('')
        # para += (line + ' ')
        
        for i in range(len(line)):
            run = docPara.add_run(line[i])
            # run = run.run
            if i >=0 and i <= bold:
                run.bold = True
                if line[i] == '*':
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                elif i in indexes:
                    run.underline = True

            elif i <= italic:
                run.italic = True

            if line[i] == '*':
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            elif i in indexes:
                run.underline = True
        docPara = doc.add_paragraph('')
doc.save('result.docx')