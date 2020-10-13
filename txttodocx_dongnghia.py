from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import RGBColor
import json

def searchForUnderline(para, text, start, end, indexes, bold, italic):
    cuts = []
    for index in indexes:
        if start <= index and index < end:
            cuts.append(index - start)

    low = 0
    for index in cuts:
        high = index
        # if high > 0:
        for i in range(low, high):
            run = para.add_run(text[i])
            run.bold = bold
            run.italic = italic
            if text[i] == '*':
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        run = para.add_run(text[index])
        run.bold = bold
        run.italic = italic
        run.underline = True
        if text[index] == '*':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        low = high + 1

    for i in range(low, end - start):
        run = para.add_run(text[i])
        run.bold = bold
        run.italic = italic
        if text[i] == '*':
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


inputsDir = 'results'
path = os.listdir(inputsDir)
doc = Document()
para = ''

docPara = doc.add_paragraph('')
docPara.add_run('KÍ HIỆU:').bold = True

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('*').font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
docPara.add_run(': các kí tự lạ nằm ngoài bảng chữ cái Tiếng Việt')

docPara = doc.add_paragraph('', style = 'ListBullet')
docPara.add_run('_').bold = True
docPara.add_run(': các chữ cái có xác suất dự đoán sai cao')

docPara = doc.add_paragraph('')

for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    # par = doc.add_paragraph(txtFile)
    contents = txtFile.split('\n')
    cumTu = False
    vidu = False
    first = False
    for line in contents:
        if line == '':
            continue

        indexes = []
        if line[line.rfind('['):line.rfind(']') + 1] != '':
            indexes = json.loads(line[line.rfind('['):line.rfind(']') + 1]) # string to list
            indexes = sorted(indexes)
        
        line = line[:line.rfind('[')]

        if line.find(' - ') != -1:
            cumTu = False
            vidu = False
            docPara = doc.add_paragraph('')
            if line.lower().find('ví dụ') != -1:
                searchForUnderline(docPara, line[:line.find(' - ')], 0, line.find(' - '), indexes, True, False)
                searchForUnderline(docPara, line[line.find(' - '):line.lower().find('ví dụ')], line.find(' - '), line.lower().find('ví dụ'), indexes, False, False)
                i = 0
                strip = line
                while strip.lower().find('ví dụ') != -1 and strip.find(':') != -1:
                    # print(strip)
                    # print(strip.lower().find('ví dụ'))
                    searchForUnderline(docPara, line[i + strip.lower().find('ví dụ'):i + strip.find(':') + 1], i + strip.lower().find('ví dụ'), i + strip.find(':') + 1, indexes, False, True)
                    
                    i += strip.find(':') + 1
                    strip = strip[strip.find(':') + 1:]

                    if strip.lower().find('ví dụ') != -1 and strip.find(':') != -1:
                        searchForUnderline(docPara, line[i:i + strip.lower().find('ví dụ')], i, i + strip.lower().find('ví dụ'), indexes, False, False)
                    else:
                        searchForUnderline(docPara, line[i:], i, len(line), indexes, False, False)
            else:
                searchForUnderline(docPara, line[:line.find(' - ')], 0, line.find(' - '), indexes, True, False)
                searchForUnderline(docPara, line[line.find(' - '):], line.find(' - '), len(line), indexes, False, False)

        elif line.lower().find('ví dụ') != -1:
            searchForUnderline(docPara, line[:line.find(':')], 0, line.find(':'), indexes, False, True)
            searchForUnderline(docPara, line[line.find(':'):], line.find(':'), len(line), indexes, False, False)

        elif cumTu == True and not line.isupper():
            if not first:
                docPara.add_run(' ').bold = True
                searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
            else:
                docPara = doc.add_paragraph('')
                searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                docParaFormat = docPara.paragraph_format
                docParaFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first = False
                

        elif line.isupper():
            if cumTu == True:
                docPara.add_run(' ').bold = True
                searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
            else:
                docPara = doc.add_paragraph('')
                searchForUnderline(docPara, line, 0, len(line), indexes, True, False)
                docParaFormat = docPara.paragraph_format
                docParaFormat.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cumTu = True
            first = True
        
        else:
            docPara.add_run(' ')
            searchForUnderline(docPara, line, 0, len(line), indexes, False, False)

doc.save('result.docx')