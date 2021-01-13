from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Inches

inputsDir = 'texts/Tu dien 005/results'
path = os.listdir(inputsDir)
# x= path[0]
# print(37 + min(x[37:].find('-') if x[37:].find('-') != -1 else 2000, x[37:].find('.')))
# print(path[0][x[37: 37 + min(x[37:].find('-') if x[37:].find('-') != -1 else 2000, x[37:].find('.'))]])
def cal(x):
    # print(int(x[33:36]) * 1000 + (int(x[37:37+min(x[37:].find('-') if x[37:].find('-') != -1 else 2000, x[37:].find('.'))])) * 10, x)
    return int(x[33:36]) * 1000 + (int(x[37:37+min(x[37:].find('-') if x[37:].find('-') != -1 else 2000, x[37:].find('.'))])) * 10
path = sorted(path, key = lambda x: cal(x))

# print(path)
doc = Document()
para = ''
docPara = doc.add_paragraph('')

# # Tags và Notes dùng cho từ điển 001 (Hoàng Phê):
# tags = {' d.': ' danh từ', 'đg.': 'động từ', ' t.': ' tính từ', 'đ.': 'đại từ', ' p.': ' phụ từ', 'k.': 'kết từ', 'tr.': 'trợ từ', ' c.': ' cảm từ', 'hoài nghi đt.': 'hoài nghi động từ'}

# notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'}

# Tags và Notes dùng cho từ điển 002 (Ng Kim Than):
tags = { 'cd.': 'ca dao', 'dt.': 'danh từ', 'đt.': 'động từ', 'gt.': 'giới từ', 'id.': 'ít dùng', 'lt.': 'liên từ', ' ng.': ' nghĩa', 'pt.': 'phụ từ', 'tht.': 'thán từ', 'tng.': 'tục ngữ', 'trt.': 'trợ từ', 'vt.': 'vị từ'}

notes = {'(id.).': '(ít dùng)', '(kng.)': '(khẩu ngữ)' , '(thgt.)': '(thông tục)', '(ph.)': '(phương ngữ)', '(vchg.)': '(văn chương)',  '(trtr.)': '(trang trọng)', '(kc.)': '(kiểu cách)', '(chm.)': '(chuyên môn)'} # còn nữa

# tags = {' dt.': ' danh từ', 'đgt.': 'động từ', ' tt.': ' tính từ', ' pht.': ' phụ từ',}

# notes = {'(id.).': '(ít dùng)', '(kng.).': '(khẩu ngữ)', '(ph.).': '(phương ngữ)', '(vch.).': '(văn chương)'}

alphabets = {'a': ['a', 'ă', 'â', 'à', 'á', 'ã', 'ả', 'ạ', 'ắ', 'ằ', 'ẵ', 'ẳ', 'ặ', 'ấ', 'ầ', 'ẫ', 'ẩ', 'ậ'], 'b': ['b'], 'c': ['c'], 'd': ['d', 'đ'], 'e': ['e', 'ê', 'é', 'è', 'ẽ', 'ẻ', 'ẹ', 'ế', 'ề', 'ễ', 'ể', 'ệ'], 'f': ['f'], 'g': ['g'], 'h': ['h'], 'i': ['i', 'í', 'ì', 'ĩ', 'ỉ', 'ị'], 'j': ['j'], 'k': ['k'], 'l': ['l'], 'm': ['m'], 'n': ['n'], 'o': ['o', 'ô', 'ơ', 'ó', 'ò', 'õ', 'ỏ', 'ọ', 'ố', 'ồ', 'ỗ', 'ổ', 'ộ', 'ớ', 'ờ', 'ỡ', 'ở', 'ợ'], 'p': ['p'], 'q': ['q'], 'r': ['r'], 's': ['s'], 't': ['t'], 'u': ['u', 'ư', 'ú', 'ù', 'ũ', 'ủ', 'ụ', 'ứ', 'ừ', 'ữ', 'ử', 'ự'], 'v': ['v'], 'w': ['w'], 'x': ['x'], 'y': ['y'], 'z': ['z']}

extrasBeforeTag = {'cn.': '(cũng nói)', 'cv.': '(cũng viết)'}  # tiếp theo sẽ là 1 từ in nghiêng có chấm cuối câu 

extrasAfterTag = {' x.': ' Xem', 'Như': 'Như'}
# newEntry = True
currentAlphabet = '`'
lastLine = '.'
pages = []
lastWord = ''
lastType = ''
lastTypeKey = ''

def hasSubString(text, tags):
    for tag in tags.keys():
        if tag in text:
            return True
    return False

def hasExtraTag(text):
    for tag in extrasAfterTag.keys():
        if tag in text:
            index = text.find(tag)
            if index > 1 and text[index - 1] != '.' or text[index - 1] != '!' or text[index - 1] != '?' or text[index - 1] != ',':
                return True

    return False

def checkFirstWord(text, currAlphabet, alphabetsDict):
    for char in alphabetsDict[currAlphabet]:
        if text.lower().startswith(char) or text.lower().startswith('"' + char):
            return True
    return False

def checkOrderNumber(text):
    for ordNum in ['IV', 'III', 'II']:
        idx = text.find(ordNum)
        if idx != - 1:
            return ordNum, idx
    return None

def checkNumber(text: str):
    for num in ['1', '2', '3', '4', '5', '6']:
        idx = text.find(num)
        if idx != - 1:
            return num, idx
    return None

def spaceSplit(a):
    if a.count(" ") == 1:
        return a.split(" ")[0]
    return " ".join(a.split(" ", 2)[:2])

countNextAlphabet = 0
count = 0
for txt in path:
    print(txt)
    txtFile = open(inputsDir + '/' + txt, encoding='utf-8', errors='ignore').read()
    txtFile = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+','', txtFile) # remove all non-XML-compatible characters
    txtFile = txtFile.replace('Ä', 'A')
    txtFile = txtFile.replace('ä', 'a')
    txtFile = txtFile.replace('“', '"')
    txtFile = txtFile.replace('”', '"')
    txtFile = txtFile.replace('¿', 't')
    txtFile = txtFile.replace(':.', 't.')
    txtFile = txtFile.replace('¡(', 't')
    
    # for key, value in notes.items():
    #     txtFile = txtFile.replace(key, value)

    # for key, value in extrasBeforeTag.items():
    #     txtFile = txtFile.replace(key, value)
        
    contents = txtFile.split('\n')
    contents = contents[:-1]    # bỏ đi kí tự đặc biệt luôn xuất hiện ở dòng cuối (đối với từ điển Nguyễn Kim Than)

    firstLine = True
    if txt[:-4].endswith("title"):
        docPara = doc.add_paragraph('')
        docPara.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for line in contents:
            if line == '':
                continue
            docPara.add_run(line).bold = True

    elif txt[:-4].endswith("center"):
        docPara = doc.add_paragraph('')
        docPara.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in contents:
            if line == '':
                continue
            docPara.add_run(line)

    else:
        docPara = doc.add_paragraph('')
        docPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for line in contents:
            if line == '':
                continue
            if firstLine:
                firstLine = False
                docPara.add_run('\t')
                docPara.add_run(line)
            
            elif lastLine.endswith(' '):
                docPara.add_run(line)
            else:
                docPara.add_run(' ' + line)
            lastLine = line
    # count += 1
    # if count == 7:
    #     break

doc.save('result.docx')
