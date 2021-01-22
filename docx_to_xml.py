from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import xml.etree.ElementTree as ET

def indent(elem, level=0):
    i = "\n" + level*"  "
    if len(elem):
        if not elem.text or not elem.text.strip():
            elem.text = i + "  "
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
        for elem in elem:
            indent(elem, level+1)
        if not elem.tail or not elem.tail.strip():
            elem.tail = i
    else:
        if level and (not elem.tail or not elem.tail.strip()):
            elem.tail = i

# fileName = 'C:/Users/Tung/Desktop/DictionaryOCR/texts/Tu dien Nguyen Kim Than/result_tu dien Nguyen Kim Than.docx'
fileName = 'result.docx'
doc = Document(fileName)
updatedDoc = Document()

tree = ET.parse('result.xml')
root = tree.getroot()

entryNo = 0
currentAlphabet = '`'
word = ''
meaning = ''
first = True
for para in doc.paragraphs:
    if len(para.runs) == 0:
        continue
    if para.runs[0].bold and not first:
        # Cắt các kí tự khoảng trắng dư thừa
        while word.startswith(' '):
            word = word[1:]
        for char in ['.', '?', ':', '!', '-', '(', '"', ';', ',', '1', 'I']:
            if word.endswith(char):
                word = word[:-1]
        # Cắt các kí tự khoảng trắng dư thừa
        while word.startswith(' '):
            word = word[1:]
        while word.endswith(' '):
            word = word[:-1]

        while meaning.startswith(' '):
            meaning = meaning[1:]
        while meaning.endswith(' '):
            meaning = meaning[:-1]

        # Xuất vào file xml
        element = root.makeelement('MUC_TU', {'Noi_dung': word})
        root.append(element)
        ET.SubElement(root[entryNo], 'Y_NGHIA', {'Noi_dung': meaning})
        entryNo += 1
        word = ''
        meaning = ''
    first = False
    # Lặp tìm vùng nào là mục từ, loại từ và định nghĩa từ
    for run in para.runs:
        if run.bold:
            word += run.text
        else:
            meaning += run.text
    meaning += '\n'
# Cắt các kí tự khoảng trắng dư thừa
while word.startswith(' '):
    word = word[1:]
for char in ['.', '?', ':', '!', '-', '(', '"', ';', ',', '1', 'I']:
    if word.endswith(char):
        word = word[:-1]
# Cắt các kí tự khoảng trắng dư thừa
while word.startswith(' '):
    word = word[1:]
while word.endswith(' '):
    word = word[:-1]

while meaning.startswith(' '):
    meaning = meaning[1:]
while meaning.endswith(' '):
    meaning = meaning[:-1]

# Xuất vào file xml
element = root.makeelement('MUC_TU', {'Noi_dung': word})
root.append(element)
ET.SubElement(root[entryNo], 'Y_NGHIA', {'Noi_dung': meaning})
entryNo += 1
        

    # # Xuất lại vào file word mới:
    # updatedDocPara = updatedDoc.add_paragraph('')
    # updatedDocPara.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # updatedDocPara.add_run(word).bold = True
    # updatedDocPara.add_run(' ' + type + ' ').italic = True
    # updatedDocPara.add_run(meaning)

indent(root)
tree.write('result.xml', encoding='utf-8', xml_declaration=True)
# updatedDoc.save('result_updated.docx')